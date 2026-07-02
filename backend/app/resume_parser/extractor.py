"""
File -> raw text extraction for the three supported resume formats.

Kept separate from parser.py (which turns raw text into structured fields)
so each concern is independently testable: extraction is "can we get text
out of this file at all", parsing is "can we understand what the text means".
"""
from __future__ import annotations

import io
from typing import Optional

from app.utils.helpers import clean_whitespace, setup_logger

logger = setup_logger("resume_parser.extractor")


def extract_text(filename: str, file_bytes: bytes) -> str:
    """Dispatch to the right extractor based on file extension."""
    lower = filename.lower()
    if lower.endswith(".pdf"):
        text = _extract_pdf(file_bytes)
    elif lower.endswith(".docx"):
        text = _extract_docx(file_bytes)
    elif lower.endswith(".txt"):
        text = file_bytes.decode("utf-8", errors="ignore")
    else:
        raise ValueError(f"Unsupported file type: {filename}. Use PDF, DOCX or TXT.")
    return clean_whitespace(text)


def _extract_pdf(file_bytes: bytes) -> str:
    try:
        import pdfplumber
    except ImportError as e:  # pragma: no cover
        raise RuntimeError("pdfplumber is required for PDF parsing. pip install pdfplumber") from e

    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text_parts.append(page_text)
    return "\n".join(text_parts)


def _extract_docx(file_bytes: bytes) -> str:
    try:
        import docx
    except ImportError as e:  # pragma: no cover
        raise RuntimeError("python-docx is required for DOCX parsing. pip install python-docx") from e

    document = docx.Document(io.BytesIO(file_bytes))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)
